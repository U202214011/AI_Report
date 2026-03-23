from __future__ import annotations
from typing import Dict, Any, List, Tuple, Optional, Set
import os
import re
import json
import base64
from io import BytesIO
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_ORIENTATION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from schema_config import get_dimension_alias_map, get_dimension_title_map

TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "export_templates")

# ---- markdown block parse ----
_HEADING_RE = re.compile(r'^(#{1,6})\s+(.*)$')
_UL_RE = re.compile(r'^\s*[-*]\s+(.*)$')
_OL_RE = re.compile(r'^\s*(\d+)\.\s+(.*)$')
_CODE_FENCE_RE = re.compile(r'^\s*```([a-zA-Z0-9_\-]*)\s*$')

# ---- inline markdown ----
_BOLD_RE = re.compile(r'\*\*(.+?)\*\*')
_ITALIC_RE = re.compile(r'\*(.+?)\*')
_UNDERLINE_RE = re.compile(r'__(.+?)__')

# ---- image placeholder ----
_IMAGE_PLACEHOLDER_FULL_RE = re.compile(
    r'^\s*\{\{image:([a-zA-Z0-9_\-\u4e00-\u9fa5]+)\}\}\s*$'
)
_IMAGE_PLACEHOLDER_INLINE_RE = re.compile(
    r'\{\{image:([a-zA-Z0-9_\-\u4e00-\u9fa5]+)\}\}'
)

_BULLET_CHAR = "\u2022"


def _safe_filename(name: str) -> str:
    s = re.sub(r'[\\/:*?"<>|]+', "_", name or "")
    s = s.strip() or "report"
    return s


def _cm(v: float) -> Cm:
    return Cm(float(v))


def _pt(v: float) -> Pt:
    return Pt(float(v))


# ---------------------------
# 基础：Word field / tab / paragraph helpers
# ---------------------------

def _add_field_run(paragraph, field_name: str):
    """
    在段落中插入 Word 域（如 PAGE / NUMPAGES）
    """
    run = paragraph.add_run()
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')

    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = f" {field_name} "

    fld_char_separate = OxmlElement('w:fldChar')
    fld_char_separate.set(qn('w:fldCharType'), 'separate')

    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')

    r = run._r
    r.append(fld_char_begin)
    r.append(instr_text)
    r.append(fld_char_separate)
    r.append(fld_char_end)


def _add_tab_stops(pf, tab_stops: List[Dict[str, Any]]):
    """
    tab_stops:
    [
      {"pos_cm": 8.0, "align":"left|center|right|decimal", "leader":"none|dot|hyphen|underscore"}
    ]
    """
    align_map = {
        "left": WD_TAB_ALIGNMENT.LEFT,
        "center": WD_TAB_ALIGNMENT.CENTER,
        "right": WD_TAB_ALIGNMENT.RIGHT,
        "decimal": WD_TAB_ALIGNMENT.DECIMAL,
    }
    leader_map = {
        "none": WD_TAB_LEADER.SPACES,
        "dot": WD_TAB_LEADER.DOTS,
        "hyphen": WD_TAB_LEADER.HYPHENS,
        "underscore": WD_TAB_LEADER.HEAVY,
    }

    for ts in tab_stops or []:
        pos = float(ts.get("pos_cm", 0))
        if pos <= 0:
            continue
        align = align_map.get(str(ts.get("align", "left")).lower(), WD_TAB_ALIGNMENT.LEFT)
        leader = leader_map.get(str(ts.get("leader", "none")).lower(), WD_TAB_LEADER.SPACES)
        pf.tab_stops.add_tab_stop(_cm(pos), align, leader)


def _set_run_font(run, family: str, size_pt: float, bold: bool = False, italic: bool = False, underline: bool = False, color: str | None = None):
    run.font.name = family
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), family)
    rFonts.set(qn("w:hAnsi"), family)
    run.font.size = _pt(size_pt)
    run.bold = bool(bold)
    run.italic = bool(italic)
    run.underline = bool(underline)
    if color:
        try:
            c = color.lstrip("#")
            if len(c) == 6:
                r, g, b = int(c[0:2], 16), int(c[2:4], 16), int(c[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)
        except Exception:
            pass


_ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _apply_paragraph_style(paragraph, p_cfg: Dict[str, Any], base_font_size: float = 11):
    line_spacing = p_cfg.get("line_spacing", 1.5)
    before = float(p_cfg.get("space_before_pt", 4))
    after = float(p_cfg.get("space_after_pt", 6))
    indent_chars = int(p_cfg.get("first_line_indent_chars", 0))

    pf = paragraph.paragraph_format
    pf.space_before = _pt(before)
    pf.space_after = _pt(after)

    if isinstance(line_spacing, (int, float)):
        pf.line_spacing = float(line_spacing)
    elif isinstance(line_spacing, dict):
        tp = str(line_spacing.get("type", "multiple")).lower()
        val = float(line_spacing.get("value", 1.5))
        if tp == "exactly":
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = _pt(val)
        elif tp == "at_least":
            pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
            pf.line_spacing = _pt(val)
        else:
            pf.line_spacing = val
    else:
        pf.line_spacing = 1.5

    if indent_chars > 0:
        pf.first_line_indent = _pt(indent_chars * base_font_size)

    if p_cfg.get("left_indent_cm") is not None:
        pf.left_indent = _cm(float(p_cfg["left_indent_cm"]))
    if p_cfg.get("right_indent_cm") is not None:
        pf.right_indent = _cm(float(p_cfg["right_indent_cm"]))

    if p_cfg.get("hanging_indent_cm") is not None:
        pf.first_line_indent = _cm(-abs(float(p_cfg["hanging_indent_cm"])))

    alignment = p_cfg.get("alignment")
    if alignment and alignment in _ALIGNMENT_MAP:
        pf.alignment = _ALIGNMENT_MAP[alignment]

    if p_cfg.get("keep_with_next") is not None:
        pf.keep_with_next = bool(p_cfg.get("keep_with_next"))
    if p_cfg.get("keep_together") is not None:
        pf.keep_together = bool(p_cfg.get("keep_together"))
    if p_cfg.get("page_break_before") is not None:
        pf.page_break_before = bool(p_cfg.get("page_break_before"))
    if p_cfg.get("widow_control") is not None:
        pf.widow_control = bool(p_cfg.get("widow_control"))

    if isinstance(p_cfg.get("tab_stops"), list):
        _add_tab_stops(pf, p_cfg.get("tab_stops") or [])


def _split_inline_markdown(text: str) -> List[Tuple[str, bool, bool, bool]]:
    if not text:
        return [("", False, False, False)]

    parts: List[Tuple[str, bool, bool, bool]] = []
    i = 0
    while i < len(text):
        if text.startswith("**", i):
            end = text.find("**", i + 2)
            if end != -1:
                parts.append((text[i + 2:end], True, False, False))
                i = end + 2
                continue

        if text.startswith("__", i):
            end = text.find("__", i + 2)
            if end != -1:
                parts.append((text[i + 2:end], False, False, True))
                i = end + 2
                continue

        if text.startswith("*", i):
            end = text.find("*", i + 1)
            if end != -1:
                parts.append((text[i + 1:end], False, True, False))
                i = end + 1
                continue

        next_special = min(
            [p for p in [text.find("**", i), text.find("__", i), text.find("*", i)] if p != -1] or [len(text)]
        )
        parts.append((text[i:next_special], False, False, False))
        i = next_special

    return parts


def _get_para_cfg(style_key: str, cfg: Dict[str, Any]) -> Dict[str, Any]:
    para_styles = cfg.get("paragraph_styles", {})
    if style_key in para_styles:
        return para_styles[style_key]
    return cfg.get("paragraph", {})


def _add_text_paragraph(doc: Document, text: str, style_key: str, cfg: Dict[str, Any]):
    fonts = cfg.get("fonts", {})
    p_cfg = _get_para_cfg(style_key, cfg)
    font_cfg = fonts.get(style_key) or fonts.get("body") or {"family": "宋体", "size_pt": 11, "bold": False}

    _WORD_STYLE_MAP = {
        "title": "Title",
        "h1": "Heading 1",
        "h2": "Heading 2",
        "h3": "Heading 3",
        "h4": "Heading 4",
        "h5": "Heading 5",
        "h6": "Heading 6",
    }
    word_style = _WORD_STYLE_MAP.get(style_key)
    if word_style:
        try:
            p = doc.add_paragraph(style=word_style)
        except KeyError:
            p = doc.add_paragraph()
    else:
        p = doc.add_paragraph()

    _apply_paragraph_style(p, p_cfg, base_font_size=float(font_cfg.get("size_pt", 11)))

    for seg, is_bold, is_italic, is_underline in _split_inline_markdown(text):
        if seg == "":
            continue
        run = p.add_run(seg)
        _set_run_font(
            run,
            family=font_cfg.get("family", "宋体"),
            size_pt=float(font_cfg.get("size_pt", 11)),
            bold=is_bold or bool(font_cfg.get("bold", False)),
            italic=is_italic,
            underline=is_underline,
            color=font_cfg.get("color")
        )
    return p


def _parse_markdown_lines(md_text: str) -> List[Tuple[str, str, int]]:
    """
    返回 [(type, text, level)]
    type: heading|ul|ol|p|blank|code|table_row
    """
    lines = (md_text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
    blocks: List[Tuple[str, str, int]] = []

    in_code = False
    code_buf: List[str] = []

    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()

        fm = _CODE_FENCE_RE.match(raw)
        if fm:
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                blocks.append(("code", "\n".join(code_buf), 0))
                code_buf = []
            i += 1
            continue

        if in_code:
            code_buf.append(raw)
            i += 1
            continue

        if not line:
            blocks.append(("blank", "", 0))
            i += 1
            continue

        hm = _HEADING_RE.match(line)
        if hm:
            lv = len(hm.group(1))
            text = hm.group(2).strip()
            blocks.append(("heading", text, lv))
            i += 1
            continue

        if "|" in raw and raw.strip().startswith("|") and raw.strip().endswith("|"):
            blocks.append(("table_row", raw.strip(), 0))
            i += 1
            continue

        um = _UL_RE.match(raw)
        if um:
            blocks.append(("ul", um.group(1).strip(), 0))
            i += 1
            continue

        om = _OL_RE.match(raw)
        if om:
            blocks.append(("ol", om.group(2).strip(), int(om.group(1))))
            i += 1
            continue

        blocks.append(("p", raw.strip(), 0))
        i += 1

    if in_code and code_buf:
        blocks.append(("code", "\n".join(code_buf), 0))

    return blocks


def _set_page_size(section, page_size: str):
    ps = str(page_size or "A4").lower()
    if ps == "letter":
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    else:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)


def _set_page_layout(doc: Document, cfg: Dict[str, Any]):
    page = cfg.get("page", {})
    margins = page.get("margin_cm", [2.5, 2.2, 2.5, 2.2])
    if len(margins) != 4:
        margins = [2.5, 2.2, 2.5, 2.2]

    sec = doc.sections[0]
    _set_page_size(sec, page.get("size", "A4"))
    orientation = str(page.get("orientation", "portrait")).lower()
    if orientation == "landscape":
        sec.orientation = WD_ORIENTATION.LANDSCAPE
        sec.page_width, sec.page_height = sec.page_height, sec.page_width

    sec.top_margin = _cm(margins[0])
    sec.right_margin = _cm(margins[1])
    sec.bottom_margin = _cm(margins[2])
    sec.left_margin = _cm(margins[3])

    if page.get("header_distance_cm") is not None:
        sec.header_distance = _cm(float(page["header_distance_cm"]))
    if page.get("footer_distance_cm") is not None:
        sec.footer_distance = _cm(float(page["footer_distance_cm"]))

    if page.get("mirror_margins") is not None:
        sec.gutter = _cm(float(page.get("gutter_cm", 0.0))) if bool(page["mirror_margins"]) else _cm(0.0)


def _apply_header_footer(doc: Document, cfg: Dict[str, Any]):
    hf = cfg.get("header_footer", {})
    if not hf:
        return

    sec = doc.sections[0]
    fonts = cfg.get("fonts", {})
    body_font = fonts.get("body", {"family": "宋体", "size_pt": 10, "bold": False})

    header_cfg = hf.get("header", {})
    if header_cfg.get("text"):
        p = sec.header.paragraphs[0] if sec.header.paragraphs else sec.header.add_paragraph()
        p.clear()
        run = p.add_run(str(header_cfg["text"]))
        _set_run_font(run, body_font.get("family", "宋体"), float(body_font.get("size_pt", 10)), bold=False, color=body_font.get("color"))

    footer_cfg = hf.get("footer", {})
    if footer_cfg.get("show_page_number"):
        p = sec.footer.paragraphs[0] if sec.footer.paragraphs else sec.footer.add_paragraph()
        p.clear()

        prefix = str(footer_cfg.get("prefix", "第 "))
        suffix = str(footer_cfg.get("suffix", " 页"))
        total_format = bool(footer_cfg.get("show_total_pages", False))
        align = footer_cfg.get("alignment", "center")

        r1 = p.add_run(prefix)
        _set_run_font(r1, body_font.get("family", "宋体"), float(body_font.get("size_pt", 10)), color=body_font.get("color"))
        _add_field_run(p, "PAGE")

        if total_format:
            r2 = p.add_run(" / ")
            _set_run_font(r2, body_font.get("family", "宋体"), float(body_font.get("size_pt", 10)), color=body_font.get("color"))
            _add_field_run(p, "NUMPAGES")

        r3 = p.add_run(suffix)
        _set_run_font(r3, body_font.get("family", "宋体"), float(body_font.get("size_pt", 10)), color=body_font.get("color"))
        p.paragraph_format.alignment = _ALIGNMENT_MAP.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    elif footer_cfg.get("text"):
        p = sec.footer.paragraphs[0] if sec.footer.paragraphs else sec.footer.add_paragraph()
        p.clear()
        run = p.add_run(str(footer_cfg["text"]))
        _set_run_font(run, body_font.get("family", "宋体"), float(body_font.get("size_pt", 10)), bold=False, color=body_font.get("color"))
        p.paragraph_format.alignment = _ALIGNMENT_MAP.get(footer_cfg.get("alignment", "center"), WD_ALIGN_PARAGRAPH.CENTER)


def _decode_base64_image(b64: str) -> Optional[bytes]:
    if not b64:
        return None
    if "," in b64 and "base64" in b64.split(",")[0]:
        b64 = b64.split(",", 1)[1]
    try:
        return base64.b64decode(b64)
    except Exception:
        return None


def _add_image_block(doc: Document, image_bytes: bytes, cfg: Dict[str, Any], caption: str | None = None):
    img_cfg = cfg.get("image", {})
    width_cm = float(img_cfg.get("max_width_cm", 16))
    align = str(img_cfg.get("alignment", "center")).lower()

    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(BytesIO(image_bytes), width=_cm(width_cm))
    p.paragraph_format.alignment = _ALIGNMENT_MAP.get(align, WD_ALIGN_PARAGRAPH.CENTER)

    p_cfg = img_cfg.get("paragraph", {})
    if p_cfg:
        _apply_paragraph_style(p, p_cfg, base_font_size=11)

    if caption:
        cap_style_key = img_cfg.get("caption_style", "caption")
        _add_text_paragraph(doc, caption, cap_style_key, cfg)


def _parse_table_row(line: str) -> List[str]:
    t = line.strip()
    if t.startswith("|"):
        t = t[1:]
    if t.endswith("|"):
        t = t[:-1]
    return [c.strip() for c in t.split("|")]


def _is_table_separator_row(cells: List[str]) -> bool:
    for c in cells:
        cc = c.replace(":", "").replace("-", "").strip()
        if cc != "":
            return False
    return True


def _add_markdown_table(doc: Document, rows: List[List[str]], cfg: Dict[str, Any]):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncols)
    table.style = cfg.get("table", {}).get("style", "Table Grid")

    fonts = cfg.get("fonts", {})
    body_font = fonts.get("body", {"family": "宋体", "size_pt": 10, "bold": False})
    head_font = cfg.get("table", {}).get("header_font") or fonts.get("h4") or body_font

    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for cidx in range(ncols):
            txt = row[cidx] if cidx < len(row) else ""
            p = cells[cidx].paragraphs[0]
            p.clear()
            run = p.add_run(txt)
            if ridx == 0:
                _set_run_font(run, head_font.get("family", "宋体"), float(head_font.get("size_pt", 10)), bold=True, color=head_font.get("color"))
            else:
                _set_run_font(run, body_font.get("family", "宋体"), float(body_font.get("size_pt", 10)), bold=False, color=body_font.get("color"))


# ---------------------------
# 按章节注入图片占位符
# ---------------------------
_HEADING_LINE_RE = re.compile(r'^(#{1,6})\s+(.*?)\s*$')


def _normalize_text_compact(text: str) -> str:
    text = (text or "").strip()
    text = re.sub(r"\s+", "", text)
    return text


def _strip_heading_prefix(text: str) -> str:
    """
    去掉常见标题编号前缀：
    一、概览 -> 概览
    （一）国家 -> 国家
    1. 国家 -> 国家
    1、国家 -> 国家
    """
    t = _normalize_text_compact(text)
    patterns = [
        r'^[一二三四五六七八九十]+、',
        r'^第[一二三四五六七八九十]+部分',
        r'^第[一二三四五六七八九十]+章',
        r'^\([一二��四五六七八九十]+\)',
        r'^（[一二三四五六七八九十]+）',
        r'^\d+\.',
        r'^\d+、',
        r'^\(\d+\)',
        r'^（\d+）',
    ]
    for p in patterns:
        t = re.sub(p, '', t)
    return t.strip()


def _parse_headings(lines: List[str]) -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []
    for i, raw in enumerate(lines):
        m = _HEADING_LINE_RE.match((raw or "").strip())
        if not m:
            continue
        level = len(m.group(1))
        text = m.group(2).strip()
        norm = _normalize_text_compact(text)
        stripped = _strip_heading_prefix(text)
        out.append({
            "line_index": i,
            "level": level,
            "text": text,
            "norm": norm,
            "stripped_norm": stripped
        })
    return out


def _build_dimension_maps(
    selected_dimensions: Optional[List[Dict[str, Any]]]
) -> Tuple[Dict[str, str], Dict[str, List[str]]]:
    default_title_map = get_dimension_title_map(include_total=False)
    default_alias_map = get_dimension_alias_map(include_total=False)

    if not selected_dimensions:
        return dict(default_title_map), dict(default_alias_map)

    title_map: Dict[str, str] = {}
    alias_map: Dict[str, List[str]] = {}

    for item in selected_dimensions:
        if not isinstance(item, dict):
            continue

        k = str(item.get("key") or "").strip().lower()
        t = str(item.get("title") or "").strip()
        aliases = item.get("aliases") or []

        if not k:
            continue

        if not t:
            t = default_title_map.get(k, k)

        alias_list = [k, t] + [str(a).strip() for a in aliases if str(a).strip()]
        alias_list += default_alias_map.get(k, [])

        seen = set()
        cleaned = []
        for a in alias_list:
            aa = str(a).strip().lower()
            if not aa or aa in seen:
                continue
            seen.add(aa)
            cleaned.append(str(a).strip())

        title_map[k] = t
        alias_map[k] = cleaned

    return title_map, alias_map


def _find_main_sections(headings: List[Dict[str, Any]], total_lines: int) -> Dict[str, Dict[str, Any]]:
    def _is_overview_title(norm: str) -> bool:
        keywords = ["概览", "总览", "概况", "总体", "整体"]
        return any(k in norm for k in keywords)

    def _is_findings_title(norm: str) -> bool:
        keywords = ["维度", "关键发现", "维度分析", "维度拆解", "维度表现"]
        return any(k in norm for k in keywords)

    def _is_cause_title(norm: str) -> bool:
        keywords = ["原因", "归因", "成因"]
        return any(k in norm for k in keywords)

    def _is_advice_title(norm: str) -> bool:
        keywords = ["建议", "策略", "行动", "举措"]
        return any(k in norm for k in keywords)

    seq: List[Tuple[str, int, str]] = []
    used = set()

    for h in headings:
        if h["level"] != 1:
            continue

        n = h.get("stripped_norm") or h["norm"]

        if _is_overview_title(n) and "overview" not in used:
            seq.append(("overview", h["line_index"], h["text"]))
            used.add("overview")
        elif _is_findings_title(n) and "findings" not in used:
            seq.append(("findings", h["line_index"], h["text"]))
            used.add("findings")
        elif _is_cause_title(n) and "cause" not in used:
            seq.append(("cause", h["line_index"], h["text"]))
            used.add("cause")
        elif _is_advice_title(n) and "advice" not in used:
            seq.append(("advice", h["line_index"], h["text"]))
            used.add("advice")

    seq = sorted(seq, key=lambda x: x[1])

    result: Dict[str, Dict[str, Any]] = {}
    for i, (name, start, title) in enumerate(seq):
        end = total_lines - 1
        if i + 1 < len(seq):
            end = seq[i + 1][1] - 1
        result[name] = {"start": start, "end": end, "title": title}
    return result


def _find_dimension_sections_in_findings(
    headings: List[Dict[str, Any]],
    findings_start: int,
    findings_end: int,
    dim_title_map: Dict[str, str]
) -> Dict[str, Dict[str, Any]]:
    alias_map = get_dimension_alias_map(include_total=False)
    found: List[Tuple[str, int, str]] = []

    for h in headings:
        if h["level"] != 2:
            continue
        if not (findings_start <= h["line_index"] <= findings_end):
            continue

        norm_title = h.get("stripped_norm") or h["norm"]
        matched_dim = None

        for dim, aliases in alias_map.items():
            candidates = set(aliases + [dim_title_map.get(dim, dim), dim])
            for alias in candidates:
                alias_norm = _normalize_text_compact(str(alias))
                if alias_norm and alias_norm in norm_title:
                    matched_dim = dim
                    break
            if matched_dim:
                break

        if matched_dim:
            found.append((matched_dim, h["line_index"], h["text"]))

    result: Dict[str, Dict[str, Any]] = {}
    for i, (dim, start, title) in enumerate(found):
        end = findings_end
        if i + 1 < len(found):
            end = found[i + 1][1] - 1
        result[dim] = {"start": start, "end": end, "title": title}
    return result


def _chart_rank_for_insert(key: str) -> int:
    k = (key or "").lower()
    if ("bar" in k) or ("柱状" in k):
        return 1
    if ("line" in k) or ("趋势" in k) or ("折线" in k):
        return 2
    if ("pie" in k) or ("饼" in k):
        return 3
    return 9


def _infer_dim_from_key(key: str, dim_alias_map: Dict[str, List[str]]) -> Optional[str]:
    k = (key or "").lower()
    for dim, aliases in dim_alias_map.items():
        for a in aliases:
            if str(a).lower() in k:
                return dim
    return None


def _normalize_chart_kind(kind: str) -> str:
    k = str(kind or "").strip().lower()
    if k in ("bar", "column", "histogram", "柱状图", "柱状"):
        return "bar"
    if k in ("line", "trend", "折线图", "折线", "趋势"):
        return "line"
    if k in ("pie", "饼图", "饼"):
        return "pie"
    return k or "chart"


def _chart_rank_for_kind(kind: str) -> int:
    k = _normalize_chart_kind(kind)
    if k == "bar":
        return 1
    if k == "line":
        return 2
    if k == "pie":
        return 3
    return 9


def _group_keys_by_meta(
    keys: List[str],
    image_meta: Dict[str, Any] | None,
    dim_alias_map: Dict[str, List[str]]
) -> Tuple[List[str], Dict[str, List[str]], List[str]]:
    overview_keys: List[str] = []
    dim_groups: Dict[str, List[str]] = {}
    remain_keys: List[str] = []

    meta_map = image_meta or {}

    for k in keys:
        meta = meta_map.get(k) or {}
        scope = str(meta.get("scope") or "").strip().lower()
        dim = str(meta.get("dimension_key") or "").strip().lower()

        if scope == "overview" or dim == "total":
            overview_keys.append(k)
            continue

        if dim:
            dim_groups.setdefault(dim, []).append(k)
            continue

        inferred = _infer_dim_from_key(k, dim_alias_map)
        if inferred:
            dim_groups.setdefault(inferred, []).append(k)
        else:
            remain_keys.append(k)

    return overview_keys, dim_groups, remain_keys


def inject_placeholders_by_sections(
    markdown_text: str,
    images: Dict[str, str] | None,
    image_meta: Dict[str, Any] | None = None,
    selected_dimensions: Optional[List[Dict[str, Any]]] = None
) -> Tuple[str, Dict[str, Any]]:
    debug: Dict[str, Any] = {
        "main_sections": {},
        "dimension_sections": {},
        "inserted": [],
        "existing_placeholders": [],
        "unmatched_to_appendix": [],
        "input_image_keys": [],
        "final_image_keys": [],
        "selected_dimensions": selected_dimensions or [],
        "image_meta_keys": sorted(list((image_meta or {}).keys()))
    }

    if not markdown_text:
        return markdown_text or "", debug
    if not images:
        return markdown_text, debug

    lines = markdown_text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    keys = [k for k, v in (images or {}).items() if v]
    debug["input_image_keys"] = list(keys)

    existing: Set[str] = set()
    for ln in lines:
        found = re.findall(r"\{\{image:([a-zA-Z0-9_\-\u4e00-\u9fa5]+)\}\}", ln or "")
        existing.update(found)
    debug["existing_placeholders"] = sorted(existing)

    keys = [k for k in keys if k not in existing]
    debug["final_image_keys"] = list(keys)
    if not keys:
        return "\n".join(lines), debug

    dim_title_map, dim_alias_map = _build_dimension_maps(selected_dimensions)

    headings = _parse_headings(lines)
    main_sections = _find_main_sections(headings, len(lines))
    debug["main_sections"] = main_sections

    dim_sections: Dict[str, Dict[str, Any]] = {}
    findings_sec = main_sections.get("findings")
    if findings_sec:
        dim_sections = _find_dimension_sections_in_findings(
            headings=headings,
            findings_start=findings_sec["start"],
            findings_end=findings_sec["end"],
            dim_title_map=dim_title_map
        )
    debug["dimension_sections"] = dim_sections

    overview_keys, dim_groups, remain_keys = _group_keys_by_meta(
        keys=keys,
        image_meta=image_meta,
        dim_alias_map=dim_alias_map
    )

    def _sort_key(x: str):
        meta = (image_meta or {}).get(x) or {}
        chart_kind = meta.get("chart_kind")
        if chart_kind:
            return (_chart_rank_for_kind(chart_kind), x)
        return (_chart_rank_for_insert(x), x)

    overview_keys = sorted(overview_keys, key=_sort_key)
    for d in list(dim_groups.keys()):
        dim_groups[d] = sorted(dim_groups[d], key=_sort_key)

    inserts: List[Tuple[int, str, str, str]] = []
    placed: Set[str] = set()

    if overview_keys and "overview" in main_sections:
        anchor = main_sections["overview"]["start"]
        for k in overview_keys:
            inserts.append((anchor, k, f"{{{{image:{k}}}}}", "overview"))
            placed.add(k)
    else:
        remain_keys.extend(overview_keys)

    for dim, ks in dim_groups.items():
        sec = dim_sections.get(dim)
        if not sec:
            remain_keys.extend(ks)
            continue
        anchor = sec["start"]
        for k in ks:
            inserts.append((anchor, k, f"{{{{image:{k}}}}}", f"dimension:{dim}"))
            placed.add(k)

    offset = 0
    for idx, key, ph, section in sorted(inserts, key=lambda it: it[0]):
        pos = idx + 1 + offset
        lines.insert(pos, "")
        lines.insert(pos + 1, ph)
        lines.insert(pos + 2, "")
        debug["inserted"].append({
            "key": key,
            "section": section,
            "insert_after_line": idx + 1,
            "actual_placeholder_line": pos + 2
        })
        offset += 3

    true_remain = [k for k in keys if k not in placed]
    if true_remain:
        lines.append("")
        lines.append("# 附录：图表")
        appendix_header_line = len(lines)

        def _appendix_sort_key(x: str):
            meta = (image_meta or {}).get(x) or {}
            chart_kind = meta.get("chart_kind")
            if chart_kind:
                return (_chart_rank_for_kind(chart_kind), x)
            return (_chart_rank_for_insert(x), x)

        for k in sorted(true_remain, key=_appendix_sort_key):
            lines.append("")
            lines.append(f"{{{{image:{k}}}}}")
            lines.append("")
            debug["inserted"].append({
                "key": k,
                "section": "appendix",
                "insert_after_line": appendix_header_line,
                "actual_placeholder_line": len(lines) - 1
            })
            debug["unmatched_to_appendix"].append(k)

    return "\n".join(lines), debug


# 新增：用户模板目录（可选）
USER_TEMPLATE_DIR = os.path.join(TEMPLATE_DIR, "user")
os.makedirs(USER_TEMPLATE_DIR, exist_ok=True)


def list_export_templates() -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []

    def scan_dir(dir_path: str, is_user: bool = False):
        if not os.path.exists(dir_path):
            return
        for fn in os.listdir(dir_path):
            if not fn.endswith(".json"):
                continue
            file_id = os.path.splitext(fn)[0]
            path = os.path.join(dir_path, fn)
            try:
                with open(path, "r", encoding="utf-8") as f:
                    cfg = json.load(f)
                out.append({
                    "id": file_id,
                    "name": cfg.get("name") or file_id,
                    "description": cfg.get("description", ""),
                    "file": fn,
                    "is_user": is_user
                })
            except Exception:
                continue

    scan_dir(TEMPLATE_DIR, is_user=False)
    scan_dir(USER_TEMPLATE_DIR, is_user=True)

    m: Dict[str, Any] = {}
    for x in out:
        m[x["id"]] = x
    return sorted(m.values(), key=lambda x: (0 if x["is_user"] else 1, x["id"]))


def load_template_config(template_id: str | None) -> Dict[str, Any]:
    if not template_id:
        template_id = "cn_management_a4"

    user_path = os.path.join(USER_TEMPLATE_DIR, f"{template_id}.json")
    if os.path.exists(user_path):
        with open(user_path, "r", encoding="utf-8") as f:
            return json.load(f)

    path = os.path.join(TEMPLATE_DIR, f"{template_id}.json")
    if not os.path.exists(path):
        path = os.path.join(TEMPLATE_DIR, "cn_management_a4.json")

    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def save_user_template_config(cfg: Dict[str, Any]) -> Dict[str, Any]:
    tid = str(cfg.get("id") or "").strip()
    if not tid:
        raise ValueError("template_config.id 不能为空")
    safe_id = re.sub(r"[^a-zA-Z0-9_\-]+", "_", tid)
    cfg["id"] = safe_id
    if not cfg.get("name"):
        cfg["name"] = safe_id

    path = os.path.join(USER_TEMPLATE_DIR, f"{safe_id}.json")
    with open(path, "w", encoding="utf-8") as f:
        json.dump(cfg, f, ensure_ascii=False, indent=2)

    return {"id": safe_id, "path": path}


def delete_user_template_config(template_id: str) -> None:
    safe_id = re.sub(r"[^a-zA-Z0-9_\-]+", "_", (template_id or "").strip())
    if not safe_id:
        raise ValueError("template_id 不能为空")
    path = os.path.join(USER_TEMPLATE_DIR, f"{safe_id}.json")
    if not os.path.exists(path):
        raise FileNotFoundError(f"用户模板 '{safe_id}' 不存在")
    os.remove(path)


# ---------------------------
# markdown -> docx（增强版）
# ---------------------------

def render_markdown_to_docx_bytes(
    markdown_text: str,
    template_cfg: Dict[str, Any],
    report_title: str | None = None,
    images: Dict[str, str] | None = None,
    image_captions: Dict[str, str] | None = None
) -> bytes:
    doc = Document()
    _set_page_layout(doc, template_cfg)
    _apply_header_footer(doc, template_cfg)

    if report_title:
        _add_text_paragraph(doc, report_title, "title", template_cfg)

    blocks = _parse_markdown_lines(markdown_text)

    i = 0
    while i < len(blocks):
        typ, text, lv = blocks[i]

        if typ == "blank":
            doc.add_paragraph("")
            i += 1
            continue

        if typ == "table_row":
            raw_rows: List[List[str]] = []
            while i < len(blocks) and blocks[i][0] == "table_row":
                raw_rows.append(_parse_table_row(blocks[i][1]))
                i += 1

            if len(raw_rows) >= 2 and _is_table_separator_row(raw_rows[1]):
                table_rows = [raw_rows[0]] + raw_rows[2:]
            else:
                table_rows = raw_rows

            _add_markdown_table(doc, table_rows, template_cfg)
            continue

        if typ == "code":
            lines = text.split("\n")
            for line in lines:
                _add_text_paragraph(doc, line, "code", template_cfg)
            i += 1
            continue

        if typ == "p":
            full = _IMAGE_PLACEHOLDER_FULL_RE.match(text)
            if full and images:
                key = full.group(1)
                img_b64 = images.get(key)
                img_bytes = _decode_base64_image(img_b64) if img_b64 else None
                if img_bytes:
                    cap = (image_captions or {}).get(key) if image_captions else None
                    _add_image_block(doc, img_bytes, template_cfg, caption=cap)
                    i += 1
                    continue

            if images and _IMAGE_PLACEHOLDER_INLINE_RE.search(text or ""):
                last = 0
                for mm in _IMAGE_PLACEHOLDER_INLINE_RE.finditer(text):
                    start, end = mm.span()
                    key = mm.group(1)

                    before = text[last:start].strip()
                    if before:
                        _add_text_paragraph(doc, before, "body", template_cfg)

                    img_b64 = images.get(key)
                    img_bytes = _decode_base64_image(img_b64) if img_b64 else None
                    if img_bytes:
                        cap = (image_captions or {}).get(key) if image_captions else None
                        _add_image_block(doc, img_bytes, template_cfg, caption=cap)
                    else:
                        _add_text_paragraph(doc, f"{{{{image:{key}}}}}", "body", template_cfg)

                    last = end

                tail = text[last:].strip()
                if tail:
                    _add_text_paragraph(doc, tail, "body", template_cfg)
                i += 1
                continue

            _add_text_paragraph(doc, text, "body", template_cfg)
            i += 1
            continue

        if typ == "heading":
            if lv == 1:
                key = "h1"
            elif lv == 2:
                key = "h2"
            elif lv == 3:
                key = "h3"
            elif lv == 4:
                key = "h4"
            elif lv == 5:
                key = "h5"
            else:
                key = "h6"
            _add_text_paragraph(doc, text, key, template_cfg)
            i += 1
            continue

        if typ == "ul":
            try:
                p = doc.add_paragraph(style="List Bullet")
                p_cfg = _get_para_cfg("list", template_cfg)
                _apply_paragraph_style(p, p_cfg, base_font_size=11)
                font_cfg = template_cfg.get("fonts", {}).get("list") or template_cfg.get("fonts", {}).get("body") or {"family": "宋体", "size_pt": 11}
                for seg, is_bold, is_italic, is_underline in _split_inline_markdown(text):
                    if not seg:
                        continue
                    run = p.add_run(seg)
                    _set_run_font(
                        run,
                        font_cfg.get("family", "宋体"),
                        float(font_cfg.get("size_pt", 11)),
                        bold=is_bold or bool(font_cfg.get("bold", False)),
                        italic=is_italic,
                        underline=is_underline,
                        color=font_cfg.get("color")
                    )
            except KeyError:
                _add_text_paragraph(doc, f"{_BULLET_CHAR} {text}", "list", template_cfg)
            i += 1
            continue

        if typ == "ol":
            try:
                p = doc.add_paragraph(style="List Number")
                p_cfg = _get_para_cfg("list", template_cfg)
                _apply_paragraph_style(p, p_cfg, base_font_size=11)
                font_cfg = template_cfg.get("fonts", {}).get("list") or template_cfg.get("fonts", {}).get("body") or {"family": "宋体", "size_pt": 11}
                for seg, is_bold, is_italic, is_underline in _split_inline_markdown(text):
                    if not seg:
                        continue
                    run = p.add_run(seg)
                    _set_run_font(
                        run,
                        font_cfg.get("family", "宋体"),
                        float(font_cfg.get("size_pt", 11)),
                        bold=is_bold or bool(font_cfg.get("bold", False)),
                        italic=is_italic,
                        underline=is_underline,
                        color=font_cfg.get("color")
                    )
            except KeyError:
                _add_text_paragraph(doc, text, "list", template_cfg)
            i += 1
            continue

        _add_text_paragraph(doc, text, "body", template_cfg)
        i += 1

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_export_filename(prefix: str = "报告", ext: str = "docx") -> str:
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{_safe_filename(prefix)}_{ts}.{ext}"